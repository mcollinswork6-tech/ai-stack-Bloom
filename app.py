import streamlit as st
import time
from query_pinecone import get_query_embedding, index 
from google import genai
from groq import Groq
from tavily import TavilyClient
from docx import Document 
from io import BytesIO
from dotenv import load_dotenv
import os

load_dotenv()

# --- 1. SETUP & CLIENTS ---
st.set_page_config(page_title="Bloom Health AI", page_icon="🌸", layout="centered")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

@st.cache_resource
def get_clients():
    return {
        "gemini": genai.Client(api_key=os.getenv(GEMINI_API_KEY)),
        "groq": Groq(api_key=os.getenv(GROQ_API_KEY)),
        "tavily": TavilyClient(api_key=os.getenv(TAVILY_API_KEY))
    }

clients = get_clients()

# --- 2. HELPER FUNCTIONS ---

def generate_report(messages):
    doc = Document()
    doc.add_heading('Bloom Health: Evidence-to-Impact Report', 0)
    for msg in messages:
        role = "USER" if msg["role"] == "user" else "BLOOM AI"
        doc.add_heading(role, level=1)
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("_" * 20)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def ask_bloom_ai(question):
    # A. Search Pinecone
    query_vector = get_query_embedding(question)
    search_results = index.query(vector=query_vector, top_k=3, include_metadata=True)
    
    best_score = search_results['matches'][0]['score'] if search_results['matches'] else 0
    context_text = ""
    detailed_sources = [] 
    
    if best_score >= 0.7:
        status_tag = f"✅ **Validated by Bloom Health Evidence Base.** (Confidence: {best_score:.1%})"
        for match in search_results['matches']:
            text_chunk = match['metadata']['text'][:2000] 
            context_text += f"\n---\n{text_chunk}"
            
            # Generate a quick 1-sentence gist of the source
            try:
                summary_req = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": f"Summarize this in 10 words: {text_chunk[:400]}"}],
                    max_tokens=25
                )
                gist = summary_req.choices[0].message.content.strip()
            except:
                gist = "Relevant maternal health documentation."

            detailed_sources.append({
                "title": match['metadata']['source'],
                "summary": gist
            })
    else:
        status_tag = "🌍 **Note: Retrieved via Real-Time Web Search.**"
        web_results = tavily.search(query=question, search_depth="basic", max_results=2)
        for res in web_results['results']:
            context_text += f"\n---\n{res['content'][:1500]}"
            detailed_sources.append({
                "title": res['url'],
                "summary": res.get('content', 'Web resource')[:80] + "..."
            })

    system_prompt = f"Maternal Health Expert. Start with: {status_tag}. Use context ONLY."

    try:
        chat_completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Context: {context_text}\n\nQ: {question}"}
            ],
            temperature=0.2,
            max_tokens=1000 
        )
        return chat_completion.choices[0].message.content, detailed_sources
    except Exception as e:
        return f"⚠️ Error: {e}", []

# --- 3. SIDEBAR & SESSION STATE ---

if "messages" not in st.session_state:
    st.session_state.messages = []
if "used_sources" not in st.session_state:
    st.session_state.used_sources = set()

side_prompt = None

with st.sidebar:
    st.title("🌸 Bloom Tools")
    
    # Source Counter Metric
    st.metric("Evidence Documents Used", len(st.session_state.used_sources))
    
    st.subheader("💡 Suggested Questions")
    suggestions = [
        "What are the warning signs of postpartum depression?",
        "How many prenatal visits are covered by CHIP Perinatal?",
        "What is the recommended weight gain during pregnancy?",
        "Tell me about the 2026 Texas Maternal Health Plan."
    ]
    for i, quest in enumerate(suggestions):
        if st.button(quest, key=f"side_sug_{i}", use_container_width=True):
            side_prompt = quest

    st.divider()

    if st.button("🗑️ Clear Chat History", use_container_width=True):
        st.session_state.messages = []
        st.session_state.used_sources = set()
    
    st.divider()

    if st.session_state.messages:
        st.subheader("Export Report")
        try:
            report_data = generate_report(st.session_state.messages)
            st.download_button(
                label="📄 Download Report (.docx)",
                data=report_data,
                file_name="Bloom_Health_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except:
            st.error("Could not generate report.")

# --- 4. MAIN UI ---

st.title("🌸 Bloom Health: Evidence-to-Impact")
st.markdown("Texas Maternal Health Guidelines and Reports.")

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# --- 5. EXECUTION LOGIC ---

user_input = st.chat_input("Ask Bloom about maternal health guidelines...")
final_prompt = side_prompt if side_prompt else user_input

if final_prompt:
    st.session_state.messages.append({"role": "user", "content": final_prompt})
    with st.chat_message("user"):
        st.markdown(final_prompt)

    with st.chat_message("assistant"):
        with st.spinner("Analyzing data and validating evidence..."):
            try:
                answer, sources = ask_bloom_ai(final_prompt)
                st.markdown(answer)
                
                if sources:
                    with st.expander("📚 Sources & Evidence"):
                        for s in sources:
                            # Update session state for unique source counter
                            if not s['title'].startswith("http"):
                                st.session_state.used_sources.add(s['title'])
                            
                            # Display formatted source
                            if s['title'].startswith("http"):
                                st.markdown(f"**[Web Resource]({s['title']})**")
                            else:
                                st.markdown(f"**Document: {s['title']}**")
                            st.caption(f"Context: {s['summary']}")
                            st.divider()
                
                st.session_state.messages.append({"role": "assistant", "content": answer})

                
            except Exception as e:
                st.error(f"Something went wrong: {e}")